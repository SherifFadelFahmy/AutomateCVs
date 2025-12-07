# generate_cv_docx.py
import json
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt

BASE_DIR = Path(__file__).parent
PROFILE_YAML = BASE_DIR / "profile.yaml"
PUB_JSON = BASE_DIR / "data" / "publications.json"
OUTPUT_DIR = BASE_DIR / "output"


def load_profile():
    with open(PROFILE_YAML, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_publications():
    if PUB_JSON.exists():
        with open(PUB_JSON, "r", encoding="utf-8") as f:
            return json.load(f)
    else:
        print("Warning: publications.json not found; run fetch_publications.py first.")
        return []


def set_heading_style(run, size=16, bold=True):
    font = run.font
    font.size = Pt(size)
    font.bold = bold


def main():
    profile = load_profile()
    publications = load_publications()
    publications.sort(key=lambda p: (p.get("year") or 0, p.get("title") or ""), reverse=True)

    max_pubs = profile.get("config", {}).get("short_cv", {}).get("max_publications", 8)
    pubs_used = publications[:max_pubs]

    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document()

    # Name
    p = doc.add_paragraph()
    run = p.add_run(profile["name"])
    set_heading_style(run, size=18, bold=True)

    # Title & contact
    doc.add_paragraph(profile["contact"]["title"])
    doc.add_paragraph(f"{profile['contact']['department']}, {profile['contact']['institution']}")
    doc.add_paragraph(f"{profile['contact']['campus']}, {profile['contact']['city_country']}")
    doc.add_paragraph(f"Phone: {profile['contact']['phone']}")
    doc.add_paragraph(f"Email: {profile['contact']['email']}")
    doc.add_paragraph(f"Website: {profile['contact']['website']}")

    # Summary
    doc.add_heading("Profile", level=1)
    doc.add_paragraph(profile.get("summary", ""))

    # Education
    doc.add_heading("Education", level=1)
    for edu in profile["education"]:
        doc.add_paragraph(
            f"{edu['degree']} — {edu['institution']}, {edu['location']} ({edu['year']})",
            style="List Bullet",
        )

    # Positions
    doc.add_heading("Academic Positions", level=1)
    doc.add_paragraph("Current:", style="List Bullet")
    for pos in profile["positions"]["current"]:
        doc.add_paragraph(f"{pos['title']} ({pos['years']})", style="List Bullet 2")

    doc.add_paragraph("Previous:", style="List Bullet")
    for pos in profile["positions"]["previous"]:
        doc.add_paragraph(f"{pos['title']} ({pos['years']})", style="List Bullet 2")

    # Research Interests
    doc.add_heading("Research Interests", level=1)
    for r in profile["research_interests"]:
        doc.add_paragraph(r, style="List Bullet")

    # Teaching
    doc.add_heading("Teaching", level=1)
    doc.add_paragraph("Undergraduate:", style="List Bullet")
    for c in profile["teaching"]["undergraduate"]:
        doc.add_paragraph(c, style="List Bullet 2")

    if profile["teaching"].get("postgraduate"):
        doc.add_paragraph("Postgraduate:", style="List Bullet")
        for c in profile["teaching"]["postgraduate"]:
            doc.add_paragraph(c, style="List Bullet 2")

    # Service
    doc.add_heading("Service and Administration", level=1)
    doc.add_paragraph("Committees and Service:", style="List Bullet")
    for s in profile["service"]["committees"]:
        doc.add_paragraph(s, style="List Bullet 2")

    doc.add_paragraph("Administrative Roles:", style="List Bullet")
    for s in profile["service"]["administration"]:
        doc.add_paragraph(s, style="List Bullet 2")

    # Awards
    awards = profile.get("awards", [])
    if awards:
        doc.add_heading("Awards and Honors", level=1)
        for a in awards:
            doc.add_paragraph(f"{a['name']} ({a['year']})", style="List Bullet")

    # Selected Publications
    doc.add_heading("Selected Publications", level=1)
    if pubs_used:
        for pub in pubs_used:
            year = pub.get("year") or "n.d."
            line = f"{pub['authors']} ({year}). {pub['title']}. {pub['venue']}."
            doc.add_paragraph(line, style="List Number")
    else:
        doc.add_paragraph("No publications available.")

    # Skills
    skills = profile.get("skills", {})
    if skills:
        doc.add_heading("Skills", level=1)
        if skills.get("technical"):
            doc.add_paragraph("Technical:", style="List Bullet")
            for s in skills["technical"]:
                doc.add_paragraph(s, style="List Bullet 2")
        if skills.get("languages"):
            doc.add_paragraph("Languages:", style="List Bullet")
            for s in skills["languages"]:
                doc.add_paragraph(s, style="List Bullet 2")

    out_path = OUTPUT_DIR / "cv_short.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
